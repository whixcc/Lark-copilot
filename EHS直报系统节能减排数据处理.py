import pandas as pd
import os
import glob
from docx import Document
from openpyxl import load_workbook
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import tkinter as tk
from tkinter import filedialog, messagebox

class App:
    def __init__(self, root):
        self.root = root
        self.root.title("节能减排数据处理")

        # 下载路径
        self.downloads_path_label = tk.Label(root, text="下载路径:")
        self.downloads_path_label.grid(row=0, column=0, padx=5, pady=5, sticky="e")
        self.downloads_path_entry = tk.Entry(root, width=50)
        self.downloads_path_entry.grid(row=0, column=1, padx=5, pady=5, sticky="ew")
        self.downloads_path_button = tk.Button(root, text="选择路径", command=self.browse_downloads_path)
        self.downloads_path_button.grid(row=0, column=2, padx=5, pady=5)

        # 输出路径
        self.output_dir_label = tk.Label(root, text="输出路径:")
        self.output_dir_label.grid(row=1, column=0, padx=5, pady=5, sticky="e")
        self.output_dir_entry = tk.Entry(root, width=50)
        self.output_dir_entry.grid(row=1, column=1, padx=5, pady=5, sticky="ew")
        self.output_dir_button = tk.Button(root, text="选择路径", command=self.browse_output_dir)
        self.output_dir_button.grid(row=1, column=2, padx=5, pady=5)

        # Excel输出路径（只读）
        self.excel_output_label = tk.Label(root, text="Excel输出路径:")
        self.excel_output_label.grid(row=2, column=0, padx=5, pady=5, sticky="e")
        self.excel_output_value = tk.StringVar()
        self.excel_output_entry = tk.Entry(root, textvariable=self.excel_output_value, state="readonly", width=60)
        self.excel_output_entry.grid(row=2, column=1, columnspan=2, padx=5, pady=5, sticky="ew")

        # 处理按钮
        self.process_button = tk.Button(root, text="处理数据", command=self.process_data)
        self.process_button.grid(row=3, column=0, columnspan=3, pady=20)

        # 设置列的权重，使Entry可以扩展
        root.grid_columnconfigure(1, weight=1)

        # 初始化路径
        self.downloads_path_entry.insert(0, r"C:\Users\cuihangming1\Downloads")
        self.output_dir_entry.insert(0, r"D:\AutoWork\直报系统")
        self.update_excel_output_path()

    def browse_downloads_path(self):
        folder_selected = filedialog.askdirectory()
        if folder_selected:
            self.downloads_path_entry.delete(0, tk.END)
            self.downloads_path_entry.insert(0, folder_selected)
            self.update_excel_output_path()

    def browse_output_dir(self):
        folder_selected = filedialog.askdirectory()
        if folder_selected:
            self.output_dir_entry.delete(0, tk.END)
            self.output_dir_entry.insert(0, folder_selected)
            self.update_excel_output_path()

    def update_excel_output_path(self):
        output_dir = self.output_dir_entry.get()
        if output_dir:
            excel_output = os.path.join(output_dir, "节能减排导入数据.xlsx")
            self.excel_output_value.set(excel_output)

    def process_data(self):
        downloads_path = self.downloads_path_entry.get()
        output_dir = self.output_dir_entry.get()
        excel_output = self.excel_output_value.get()

        try:
            self.process_excel(downloads_path, output_dir, excel_output)
            self.update_word_table(output_dir, excel_output)
            messagebox.showinfo("完成", "直报系统中节能减排数据已替换")
        except Exception as e:
            messagebox.showerror("错误", f"发生错误: {str(e)}")

    def process_excel(self, downloads_path, output_dir, excel_output):
        # 查找最新的Excel文件
        pattern = os.path.join(downloads_path, "*非工业其他行业节能减排监测统计表*.xls*")
        files = glob.glob(pattern)
        if not files:
            raise Exception("未找到符合条件的Excel文件")
        latest_file = max(files, key=os.path.getctime)

        # 需要保留的行（按指定顺序）
        keep_rows = [
            "能源消费总量", "综合能源消费量", "耗电量", "汽油消耗量",
            "柴油消耗量", "天然气消耗量", "热力消耗量", "其他能源消耗量",
            "营业收入(可比价)", "增加值(可比价)", "万元营业收入综合能耗(可比价)",
            "万元增加值综合能耗(可比价)", "二氧化碳排放总量"
        ]

        # 读取和处理Excel
        df = pd.read_excel(latest_file)
        row_order = {row: index for index, row in enumerate(keep_rows)}
        df_filtered = df[df.iloc[:, 0].isin(keep_rows)].copy()
        df_filtered['sort_key'] = df_filtered.iloc[:, 0].map(row_order)
        df_filtered = df_filtered.sort_values('sort_key').drop('sort_key', axis=1)

        # 选择和重命名列
        columns_to_keep = df_filtered.columns[[0, 5, 3, 6]]  # 修改这里，选择正确的列索引
        df_final = df_filtered[columns_to_keep]
        df_final.columns = ['指标', '上年同期累计值', '当期年累值', '同比变化%']

        # 创建输出目录并保存
        os.makedirs(output_dir, exist_ok=True)
        df_final.to_excel(excel_output, index=False)
        print(f"Excel文件已处理完成，保存至: {excel_output}")

    def format_number(self, value):
        try:
            num = float(value)
            return f"{num:.2f}"
        except (ValueError, TypeError):
            return value

    def format_percentage(self, value):
        try:
            num = float(value)
            if num > 0:
                return f"增加{abs(num):.2f}%"
            elif num < 0:
                return f"减少{abs(num):.2f}%"
            else:
                return "持平"
        except (ValueError, TypeError):
            return value

    def update_word_table(self, output_dir, excel_output):
        # 读取处理后的Excel数据
        wb = load_workbook(excel_output)
        sheet = wb.active

        # 获取Excel数据
        excel_data = []
        for row in sheet.iter_rows(min_row=2):
            row_data = []
            for col in [sheet.cell(row[0].row, sheet.max_column - 2),  # 上年同期累计值
                        sheet.cell(row[0].row, sheet.max_column - 1),  # 当期年累值
                        sheet.cell(row[0].row, sheet.max_column)]:  # 同比变化%
                value = col.value
                if any(keyword in str(row[0].value) for keyword in
                       ['营业收入(可比价)', '增加值(可比价)', '二氧化碳排放总量']):
                    value = self.format_number(value)
                row_data.append(str(value) if value is not None else "")
            excel_data.append(row_data)

        # 查找和处理Word文档
        word_files = [f for f in glob.glob(os.path.join(output_dir, "*能源节约与生态环境保护总结*.docx"))
                      if not os.path.basename(f).startswith('~')]

        if not word_files:
            raise Exception("未找到符合条件的Word文档")

        word_path = word_files[0]
        print(f"正在处理Word文件: {word_path}")

        # 更新Word表格
        doc = Document(word_path)
        if len(doc.tables) == 0:
            raise Exception("Word文档中没有表格")

        table = doc.tables[0]
        for i, row in enumerate(table.rows[1:]):  # 跳过表头
            if i < len(excel_data):
                for j, col in enumerate([4, 5, 6]):  # 第4、5、6列
                    if j < len(excel_data[i]):
                        cell = row.cells[col - 1]
                        cell.text = excel_data[i][j]

                        # 设置单元格格式
                        paragraph = cell.paragraphs[0]
                        paragraph.clear()
                        run = paragraph.add_run(excel_data[i][j])
                        run.font.name = '仿宋'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
                        run.font.size = Pt(12)
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 构建替换字典
        df = pd.read_excel(excel_output)
        indicators = {
            '【1】': ('综合能源消费量', '当期年累值'),
            '【2】': ('综合能源消费量', '同比变化%'),
            '【3】': ('耗电量', '当期年累值'),
            '【4】': ('耗电量', '同比变化%'),
            '【5】': ('营业收入(可比价)', '当期年累值'),
            '【6】': ('营业收入(可比价)', '同比变化%'),
            '【7】': ('增加值(可比价)', '当期年累值'),
            '【8】': ('增加值(可比价)', '同比变化%'),
            '【9】': ('万元营业收入综合能耗(可比价)', '当期年累值'),
            '【10】': ('万元营业收入综合能耗(可比价)', '同比变化%')
        }

        replacements = {}
        for placeholder, (indicator, column) in indicators.items():
            try:
                value = df.loc[df['指标'] == indicator, column].values[0]
                if column == '同比变化%':
                    replacement_value = self.format_percentage(value)
                else:
                    replacement_value = self.format_number(value)
                replacements[placeholder] = replacement_value
            except IndexError:
                print(f"警告：未找到指标 {indicator} 的数据")
                replacements[placeholder] = ''

        # 调试信息，打印替换字典
        print("替换字典内容：")
        for key, value in replacements.items():
            print(f"{key}: {value}")

        # 替换文档中的占位符
        def replace_placeholder(doc_obj, replacements):
            # 替换段落中的占位符
            for p in doc_obj.paragraphs:
                for placeholder, replacement in replacements.items():
                    if placeholder in p.text:
                        # 创建一个新的段落，避免占位符被拆分
                        inline = p.runs
                        # 合并 runs
                        text = ''.join([run.text for run in inline])
                        # 替换文本
                        text = text.replace(placeholder, replacement)
                        # 清空原有 runs
                        for idx in range(len(inline)-1, -1, -1):
                            p._element.remove(inline[idx]._element)
                        # 添加新的 run
                        run = p.add_run(text)
                        run.font.name = '仿宋'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
                        run.font.size = Pt(12)
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT  # 根据需要设置对齐方式

            # 替换表格中的占位符
            for table in doc_obj.tables:
                for row in table.rows:
                    for cell in row.cells:
                        replace_placeholder(cell, replacements)

        replace_placeholder(doc, replacements)

        doc.save(word_path)
        print("Word文档更新完成")

def main():
    root = tk.Tk()
    app = App(root)
    root.mainloop()

if __name__ == "__main__":
    main()
